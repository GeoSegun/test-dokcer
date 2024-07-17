import pandas as pd
import streamlit as st
from docx import Document
from PIL import Image
import base64
import io
from collections import Counter
import random
from sqlalchemy import create_engine

from dotenv import load_dotenv
import os

# Load environment variables
load_dotenv()

col1, col2 = st.columns((2))

st.toast('I am happy to see you Here!!!', icon='😍')
image = Image.open('cybersoc.JPG')

st.image(image, use_column_width=True)
st.title("MONTHLY TICKET SUMMARY")

with st.expander("**ABOUT**"):
    st.markdown("""
        This web app helps to reduce the work time in providing the ticket summary for :globe_with_meridians: **CyberLab** :shield:
        * 🔹 **Python 🐍 libraries:** base64, pandas, streamlit, openpyxl 📈🔍, xlsxwriter, docx, plotly, random
        * 🔹 **Data source:** Uploaded from Local Machine  🚀🔍


        """)

# Load Data
def load_data(file_path):
    path = file_path
    excel = pd.read_excel(path)
    
    # Check if 'Legend:' exists in the 'Subject' column
    if 'Legend:' in excel['Subject'].values:
        legend_index = excel[excel['Subject'] == 'Legend:'].index[0]
        excel = excel.iloc[:legend_index]
    else:
        # If 'Legend:' is not found, use the entire DataFrame
        legend_index = len(excel)
    
    selected_columns = [0, 1, 2, 3, 9, -1]
    data = excel.iloc[:, selected_columns]
    return data




# Upload Multiple Data Files
def upload_data():
    st.write("Upload your data files")
    
    data_files = st.file_uploader(
        ":file_folder: Choose data file(s) in Excel format",
        type=["xlsx"],
        accept_multiple_files=True
    )

    if data_files:
        data_frames = []
        for file in data_files:
            data_frames.append(load_data(file))

        if len(data_frames) > 0:
            merged_df = pd.concat(data_frames, ignore_index=True)

            # # Display the uploaded data files
            # st.write("Uploaded Data Files:")
            # for i, file in enumerate(data_files):
            #     st.write(f"File {i+1}: {file.name}")

            # Display the merged DataFrame in a table
            st.write("Data uploaded and merged successfully")

            return merged_df  # Return the merged DataFrame

# Call the upload_data() function to run the Streamlit app and load the data
merged_df = upload_data()

filtered_df = merged_df[merged_df['Subject'].notna()]
def filedownload(df):
    # Create Excel writer object
    towrite = io.BytesIO()
    excel_writer = pd.ExcelWriter(towrite, engine="xlsxwriter")
    df.to_excel(excel_writer, index=False, sheet_name="Sheet1")
    excel_writer.save()
    excel_data = towrite.getvalue()

    # Create download link
    b64 = base64.b64encode(excel_data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="ticket.xlsx">Download Excel File</a>'
    return href
st.markdown(filedownload(merged_df), unsafe_allow_html=True)
# convert the 'Date' column to pandas datetime format
filtered_df['Create date'] = pd.to_datetime(filtered_df['Create date'], format='%d/%m/%Y %H:%M:%S')

filtered_df['Create date'] = pd.to_datetime(filtered_df['Create date'])

startDate = pd.to_datetime(filtered_df['Create date']).min()
endDate = pd.to_datetime(filtered_df['Create date']).max()

st.write("This Web application just reduces the work time, interact with it to explore other functionality. when you select the 'Final Data', you can download the excel file of the data from the second download link")

# Custom CSS to reduce the size of the date input widgets
st.markdown("""
<style>
input[type="date"] {
  width: 150px;
  height: 30px;
}
</style>
""", unsafe_allow_html=True)

col1, col2 = st.columns((2))

with col1:
    date1= pd.to_datetime(st.date_input("startDate", startDate))

with col2:
    date2= pd.to_datetime(st.date_input("endDate", endDate))

filtered_df = filtered_df[(filtered_df["Create date"] >= date1) & (filtered_df["Create date"] <= date2)].copy()

unique_customer = filtered_df.iloc[:, 3].unique()
selected_customer = st.sidebar.selectbox("select the customer", unique_customer)

filtered_df = filtered_df[filtered_df.iloc[:, 3] == selected_customer]
# st.write(filtered_df)
data = filtered_df.copy()
#Remove "INTELLIGENCE:" prefix from the first column
data['Subject'] = data['Subject'].str.replace(r'INTELLIGENCE:\s+', '', regex=True)
#Remove text inside square brackets from the first column
data['Subject'] = data['Subject'].str.replace(r'\[.*?\]', '', regex=True)
data['Subject'] = data['Subject'].apply(lambda x: 'New Certificate Registration' if pd.notna(x) and 'certificate registration' in x.lower() else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Flash Report' if pd.notna(x) and ('Flash Report' in x or 'Flash Alert' in x or 'Flash' in x) else x)
data['Subject'] = data['Subject'].apply(lambda x: 'New Domains Registration' if pd.notna(x) and 'domain registration' in x.lower() else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Downtime monitoring' if pd.notna(x) and 'Downtime monitoring' in x else x)
data['Subject'] = data['Subject'].apply(lambda x: 'App on Third Party Store' if pd.notna(x) and 'Third Party' in x else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Suspicious Domain' if pd.notna(x) and 'Suspicious Domain' in x else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Twitter Impersonation' if pd.notna(x) and 'twitter impersonation' in x.lower() else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Telegram Impersonation' if pd.notna(x) and 'telegram' in x.lower() else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Facebook Impersonation' if pd.notna(x) and 'facebook impersonation' in x.lower() else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Youtube Impersonation' if pd.notna(x) and 'youtube impersonation' in x.lower() else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Fake Opay Investment Website' if pd.notna(x) and 'investment' in x.lower() else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Vulnerability Assessment Report' if pd.notna(x) and 'Vulnerability Assessment Report' in x else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Monthly Intelligence Report' if pd.notna(x) and 'monthly intelligence report' in x.lower() else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Monthly Patch Tuesday' if pd.notna(x) and 'monthly patch tuesday' in x.lower() else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Monthly Patch Tuesday' if pd.notna(x) and 'monthly patch tuesday' in x.lower() else x)
data['Subject'] = data['Subject'].apply(lambda x: 'Malware IOCs' if pd.notna(x) and 'malware iocs' in x.lower() else x)


unique_counts = data['Subject'].value_counts()
unique = unique_counts.reset_index()
unique = unique.rename(columns={"index": "Subject", "Subject": "Values"})
unique = unique[~unique['Subject'].str.lower().str.contains('service availability check')]

test = merged_df.copy()

test = test.drop(['Ticket link'], axis=1)

# Create a duplicate of the 'Subject' column
ticket_categories = test['Subject'].copy()

# Assign the duplicated column as the last column with the desired name
test['Ticket Categories'] = ticket_categories
#Remove "INTELLIGENCE:" prefix from the first column
test['Ticket Categories'] = test['Ticket Categories'].str.replace(r'INTELLIGENCE:\s+', '', regex=True)
#Remove text inside square brackets from the first column
test['Ticket Categories'] = test['Ticket Categories'].str.replace(r'\[.*?\]', '', regex=True)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'New Certificate Registration' if pd.notna(x) and 'certificate registration' in x.lower() else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Flash Report' if pd.notna(x) and ('Flash Report' in x or 'Flash Alert' in x or 'Flash' in x) else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'New Domains Registration' if pd.notna(x) and 'domain registration' in x.lower() else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Downtime monitoring' if pd.notna(x) and 'Downtime monitoring' in x else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'App on Third Party Store' if pd.notna(x) and 'Third Party' in x else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Suspicious Domain' if pd.notna(x) and 'Suspicious Domain' in x else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Twitter Impersonation' if pd.notna(x) and 'twitter impersonation' in x.lower() else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Telegram Impersonation' if pd.notna(x) and 'telegram' in x.lower() else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Facebook Impersonation' if pd.notna(x) and 'facebook impersonation' in x.lower() else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Youtube Impersonation' if pd.notna(x) and 'youtube impersonation' in x.lower() else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Fake Opay Investment Website' if pd.notna(x) and 'investment' in x.lower() else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Vulnerability Assessment Report' if pd.notna(x) and 'Vulnerability Assessment Report' in x else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Monthly Intelligence Report' if pd.notna(x) and 'monthly intelligence report' in x.lower() else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Monthly Patch Tuesday' if pd.notna(x) and 'monthly patch tuesday' in x.lower() else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Monthly Patch Tuesday' if pd.notna(x) and 'monthly patch tuesday' in x.lower() else x)
test['Ticket Categories'] = test['Ticket Categories'].apply(lambda x: 'Malware IOCs' if pd.notna(x) and 'malware iocs' in x.lower() else x)
test = test.dropna()

# Function to create a download link for the DataFrame
def filedownload(df):
    towrite = io.BytesIO()
    excel_writer = pd.ExcelWriter(towrite, engine='xlsxwriter')
    df.to_excel(excel_writer, index=False, sheet_name='Sheet1')
    excel_writer.save()
    excel_data = towrite.getvalue()
    b64 = base64.b64encode(excel_data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="ticket_categories.xlsx">Download Excel File</a>'
    return href

# Sidebar button to download the `test` DataFrame
if st.sidebar.button('Download Ticket Categories'):
    st.sidebar.markdown(filedownload(test), unsafe_allow_html=True)


def filedownload(df):
    # Create Excel writer object
    towrite = io.BytesIO()
    excel_writer = pd.ExcelWriter(towrite, engine="xlsxwriter")
    df.to_excel(excel_writer, index=False, sheet_name="Sheet1")
    excel_writer.save()
    excel_data = towrite.getvalue()

    # Create download link
    b64 = base64.b64encode(excel_data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="ticket.xlsx">Download Excel File</a>'
    return href
st.markdown(filedownload(unique), unsafe_allow_html=True)

data_to_display = st.sidebar.selectbox(
    "Select DataFrame to Display",
    options=["Ticket Informations", "Final Data", "Ticket Categories"]
)

if data_to_display == "Ticket Informations":
    if merged_df is not None:
        st.write(merged_df)
    else:
        st.write("No data uploaded yet. Please upload your data files.")
elif data_to_display == "Final Data":
    st.write(unique)
elif data_to_display == "Ticket Categories":
    st.write(test)

    if st.button('Send to Database'):
        try:
            # Establish connection to the database
            db_host = os.getenv('DB_HOST')
            db_port = os.getenv('DB_PORT')
            db_name = os.getenv('DB_NAME')
            db_user = os.getenv('DB_USER')
            db_password = os.getenv('DB_PASSWORD')

            engine = create_engine(f'postgresql://{db_user}:{db_password}@{db_host}:{db_port}/{db_name}')

            # Connect to the database
            with engine.connect() as connection:
                
                # Retrieve existing data from the database
                existing_data = pd.read_sql_table('ticket_categories', con=connection)

                # Filter out rows from test DataFrame that already exist in the database
                test_unique = test[~test.apply(tuple, axis=1).isin(existing_data.apply(tuple, axis=1))]

                if not test_unique.empty:
                    # Insert data into the ticket_categories table
                    test_unique.to_sql('ticket_categories', connection, if_exists='append', index=False, method='multi')
                    st.success("Data was uploaded successfully!")
                else:
                    st.warning("No new data to upload. All records are already in the database.")

        except Exception as e:
            st.error(f"Error uploading data to database: {e}")


# Create a function to generate and save the Word document
def generate_word_doc(df):
    # Create a new Word document
    doc = Document()

    # Count the occurrences of each unique subject
    subject_counts = Counter(df['Subject'])

    # Keep track of printed subjects
    printed_subjects = set()

    # Loop through the DataFrame and add information to the document
    for index, row in df.iterrows():
        subject = row['Subject']
        if "Service Availability Check" in subject:
            continue

        ticket_id = row['Ticket link']
        status = row['State']

        # Check if the subject has multiple occurrences
        if subject_counts[subject] > 1 and subject not in printed_subjects:
            # Add information with volume count
            doc.add_paragraph("SUBJECT: " + subject)
            doc.add_paragraph("Volume count: " + str(subject_counts[subject]))
            doc.add_paragraph("STATUS: " + status)
            printed_subjects.add(subject)
            doc.add_paragraph("")  # Add an empty line after each entry
        elif subject not in printed_subjects:
            # Add information without volume count
            doc.add_paragraph("TICKET ID: " + ticket_id)
            doc.add_paragraph("SUBJECT: " + subject)
            doc.add_paragraph("STATUS: " + status)
            printed_subjects.add(subject)
            doc.add_paragraph("")  # Add an empty line after each entry

    # Save the document to a file
    doc_path = "ticket_information.docx"
    doc.save(doc_path)
    return doc_path

# Call the function to generate the Word document
doc_path = generate_word_doc(filtered_df)


# Add a download button to the Streamlit app
st.sidebar.download_button(
    label="Download Ticket Information",
    data=open(doc_path, "rb").read(),
    file_name="ticket_information.docx",
    mime="application/docx"
)
